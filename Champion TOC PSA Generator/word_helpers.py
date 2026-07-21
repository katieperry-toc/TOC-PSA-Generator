"""Reusable python-docx helpers.

Every function here operates on Document/Paragraph/run/table/cell objects
and has no PSA- or Addendum-specific meaning — it doesn't know what a
"cover letter" or "signature table" is. That business meaning lives in
psa_builder.py / addendum_builder.py, which import this module.

Only scoping.clean() is imported from outside this module (a plain string
helper with no docx dependency), so this stays a low-level, reusable layer.
"""

import re
from copy import deepcopy
from typing import Dict, Iterable, List, Sequence, Tuple

from docx import Document
from docx.image.image import Image as DocxImage
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from config import BLACK
from scoping import clean


def iter_table_paragraphs(table) -> Iterable:
    seen = set()
    for row in table.rows:
        for cell in row.cells:
            cell_id = id(cell._tc)
            if cell_id in seen:
                continue
            seen.add(cell_id)
            yield from cell.paragraphs
            for nested_table in cell.tables:
                yield from iter_table_paragraphs(nested_table)

def iter_all_paragraphs(doc: Document) -> Iterable:
    yield from doc.paragraphs
    for table in doc.tables:
        yield from iter_table_paragraphs(table)
    for section in doc.sections:
        for area in (section.header, section.footer):
            yield from area.paragraphs
            for table in area.tables:
                yield from iter_table_paragraphs(table)

def descendants(element, tag_name: str):
    yield from element.iter(qn(tag_name))

def set_run_black(run) -> None:
    """Force a run to true black and remove highlighting."""
    run.font.color.rgb = BLACK
    run_properties = run._r.get_or_add_rPr()
    color = run_properties.find(qn("w:color"))
    if color is not None:
        color.set(qn("w:val"), "000000")
    highlight = run_properties.find(qn("w:highlight"))
    if highlight is not None:
        run_properties.remove(highlight)

def run_is_red(run) -> bool:
    """Return True for the red placeholder colors used in the PSA template."""
    color = run.font.color.rgb
    if color is None:
        color_node = run._r.get_or_add_rPr().find(qn("w:color"))
        if color_node is not None:
            value = clean(color_node.get(qn("w:val"))).upper()
            return value in {"EE0000", "FF0000", "C00000"}
        return False
    return str(color).upper() in {"EE0000", "FF0000", "C00000"}

def clean_inserted_run(run) -> None:
    """Make generated PSA text true black while preserving font and sizing."""
    set_run_black(run)

    run_properties = run._r.get_or_add_rPr()
    highlight = run_properties.find(qn("w:highlight"))
    if highlight is not None:
        run_properties.remove(highlight)

def replace_text_in_paragraph(paragraph, old: str, new: str) -> bool:
    """Replace text across runs without removing drawings or rebuilding XML."""
    if not old:
        return False

    changed = False

    while old in paragraph.text:
        runs = list(paragraph.runs)
        text_runs = [run for run in runs if not _run_contains_drawing(run)]
        full_text = "".join(run.text for run in text_runs)
        start = full_text.find(old)

        if start < 0:
            break

        end = start + len(old)
        cursor = 0
        affected = []

        for run in text_runs:
            run_start = cursor
            run_end = cursor + len(run.text)
            if run_end > start and run_start < end:
                affected.append((run, run_start, run_end))
            cursor = run_end

        if not affected:
            break

        first_run, first_start, _ = affected[0]
        last_run, last_start, _ = affected[-1]

        prefix = first_run.text[: start - first_start]
        suffix = last_run.text[end - last_start :]
        first_run.text = prefix + str(new) + suffix
        clean_inserted_run(first_run)

        for run, _, _ in affected[1:]:
            _clear_run_text_only(run)

        changed = True

    return changed

def replace_everywhere(doc: Document, replacements: Dict[str, str]) -> None:
    """
    Apply all replacements to each paragraph before moving to the next one.

    This avoids repeatedly mutating the same paragraph dozens of times,
    which helps reduce formatting drift.
    """
    cleaned = {
        str(old): "" if new is None else str(new)
        for old, new in replacements.items()
        if old
    }

    for paragraph in iter_all_paragraphs(doc):
        # Never rewrite a paragraph that owns a floating signature, image,
        # shape, or other anchored object. Even harmless-looking run edits can
        # cause Word to recalculate the object's anchor position.
        if _paragraph_contains_drawing(paragraph):
            continue

        text = "".join(run.text for run in paragraph.runs)

        if not text:
            continue

        # Skip paragraphs that don't contain any placeholders.
        if not any(old in text for old in cleaned):
            continue

        for old, new in cleaned.items():
            if old in text:
                replace_text_in_paragraph(
                    paragraph,
                    old,
                    new,
                )

def replace_visible_text_preserving_runs(paragraph, new_text: str) -> None:
    """Replace visible text without deleting drawings, anchors, or paragraph XML."""
    set_paragraph_value(paragraph, str(new_text))

def replace_regex_in_paragraph(paragraph, pattern: str, replacement: str) -> bool:
    text = paragraph.text
    updated = re.sub(pattern, replacement, text)
    if updated == text:
        return False
    set_paragraph_value(paragraph, updated)
    return True

def set_checkbox(paragraph, checked: bool) -> None:
    for checkbox in descendants(paragraph._p, "w14:checkbox"):
        checked_node = checkbox.find(qn("w14:checked"))
        if checked_node is not None:
            checked_node.set(qn("w14:val"), "1" if checked else "0")
    for content_control in descendants(paragraph._p, "w:sdt"):
        if not any(True for _ in descendants(content_control, "w14:checkbox")):
            continue
        for text_node in descendants(content_control, "w:t"):
            text_node.text = "☒" if checked else "☐"
    for run in paragraph.runs:
        set_run_black(run)

def _copy_run_format(source_run, target_run) -> None:
    """Copy the complete Word run formatting, including the template font."""
    if source_run is None:
        return
    source_rpr = source_run._r.rPr
    if source_rpr is not None:
        target_rpr = target_run._r.get_or_add_rPr()
        target_run._r.remove(target_rpr)
        target_run._r.insert(0, deepcopy(source_rpr))

def _copy_paragraph_format(source_paragraph, target_paragraph) -> None:
    """Copy paragraph spacing, indentation, alignment, and style settings."""
    source_ppr = source_paragraph._p.pPr
    if source_ppr is not None:
        target_ppr = target_paragraph._p.pPr
        if target_ppr is not None:
            target_paragraph._p.remove(target_ppr)
        target_paragraph._p.insert(0, deepcopy(source_ppr))

def _run_contains_drawing(run) -> bool:
    """Return True when a run contains an image, shape, object, or drawing anchor."""
    drawing_tags = {
        qn("w:drawing"),
        qn("w:pict"),
        qn("w:object"),
    }
    return any(node.tag in drawing_tags for node in run._r.iter())

def _paragraph_contains_drawing(paragraph) -> bool:
    """Return True when a paragraph owns a floating or inline Word object."""
    return any(_run_contains_drawing(run) for run in paragraph.runs)

def _clear_run_text_only(run) -> None:
    """Clear visible text without deleting drawings or other anchored objects."""
    for text_node in list(run._r.iter(qn("w:t"))):
        text_node.text = ""
    for tab_node in list(run._r.iter(qn("w:tab"))):
        parent = tab_node.getparent()
        if parent is not None:
            parent.remove(tab_node)
    for break_node in list(run._r.iter(qn("w:br"))):
        parent = break_node.getparent()
        if parent is not None:
            parent.remove(break_node)

def set_paragraph_value(paragraph, value: str, bold: bool | None = None) -> None:
    """Replace visible text while preserving the original paragraph and drawings.

    Floating signatures and other Word objects remain attached to their
    original runs because this function never rebuilds the paragraph XML.
    """
    value = "" if value is None else str(value)
    runs = list(paragraph.runs)
    editable_runs = [run for run in runs if not _run_contains_drawing(run)]

    if editable_runs:
        target_run = editable_runs[0]
        target_run.text = value
        clean_inserted_run(target_run)

        for run in editable_runs[1:]:
            _clear_run_text_only(run)
    else:
        target_run = paragraph.add_run(value)
        clean_inserted_run(target_run)

    if bold is not None:
        target_run.bold = bold

def set_cell_lines(cell, lines: Sequence[Tuple[str, bool]]) -> None:
    """Write cell lines while preserving each template paragraph's formatting."""
    original_tcPr = deepcopy(cell._tc.tcPr)

    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        paragraphs = [cell.add_paragraph()]

    source_paragraphs = list(paragraphs)
    fallback_paragraph = source_paragraphs[-1]

    while len(paragraphs) < len(lines):
        source = source_paragraphs[min(len(paragraphs), len(source_paragraphs) - 1)]
        new_paragraph = cell.add_paragraph()
        _copy_paragraph_format(source, new_paragraph)
        if source.runs:
            new_run = new_paragraph.add_run("")
            _copy_run_format(source.runs[0], new_run)
        paragraphs.append(new_paragraph)

    for index, (text, bold) in enumerate(lines):
        paragraph = paragraphs[index]
        source = (
            source_paragraphs[min(index, len(source_paragraphs) - 1)]
            if source_paragraphs
            else fallback_paragraph
        )
        _copy_paragraph_format(source, paragraph)

        if not paragraph.runs and source.runs:
            new_run = paragraph.add_run("")
            _copy_run_format(source.runs[0], new_run)

        set_paragraph_value(paragraph, text, bold)

    for paragraph in paragraphs[len(lines):]:
        remove_paragraph(paragraph)

    if original_tcPr is not None:
        current_tcPr = cell._tc.get_or_add_tcPr()
        cell._tc.remove(current_tcPr)
        cell._tc.insert(0, original_tcPr)

def clear_cell_text_preserving_format(cell) -> None:
    """Clear answer text while preserving table geometry and embedded objects."""
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        paragraphs = [cell.add_paragraph()]

    for paragraph in paragraphs:
        runs = list(paragraph.runs)
        editable_runs = [run for run in runs if not _run_contains_drawing(run)]

        if editable_runs:
            editable_runs[0].text = ""
            for run in editable_runs[1:]:
                _clear_run_text_only(run)
        else:
            paragraph.add_run("")

def remove_table_row(table, row) -> None:
    table._tbl.remove(row._tr)

def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)

def set_table_cell(cell, value: str, bold: bool | None = None) -> None:
    lines = str(value).splitlines() or [""]
    set_cell_lines(cell, [(line, bool(bold)) for line in lines])

def find_paragraph(doc: Document, exact_text: str):
    for paragraph in iter_all_paragraphs(doc):
        if clean(paragraph.text) == exact_text:
            return paragraph
    return None

def _set_row_minimum_height(row, points: float) -> None:
    """Set a minimum row height while allowing Word to expand for wrapped text."""
    from docx.oxml import OxmlElement

    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(int(points * 20)))
    tr_height.set(qn("w:hRule"), "atLeast")

def _image_part_dimensions(image_part) -> Tuple[int, int]:
    """Return image dimensions in pixels without relying on Pillow."""
    try:
        image = DocxImage.from_blob(image_part.blob)
        return int(image.px_width), int(image.px_height)
    except Exception:
        return 0, 0

def _drawing_image_parts(doc: Document, paragraph) -> List:
    """Return image parts referenced by drawings in a paragraph."""
    parts = []
    for blip in paragraph._p.iter(qn("a:blip")):
        rel_id = blip.get(qn("r:embed"))
        if not rel_id:
            continue
        rel = doc.part.rels.get(rel_id)
        if rel is not None and hasattr(rel, "target_part"):
            parts.append(rel.target_part)
    return parts

def _find_handwritten_signature_blob(doc: Document) -> bytes | None:
    """Find Meghan's wide handwritten-signature image in the template package."""
    best_blob = None
    best_ratio = 0.0

    for part in doc.part.package.parts:
        content_type = clean(getattr(part, "content_type", ""))
        if not content_type.startswith("image/"):
            continue

        width_px, height_px = _image_part_dimensions(part)
        if width_px <= 0 or height_px <= 0:
            continue

        ratio = width_px / height_px
        # Meghan's signature is the only very wide, short image in the template.
        if ratio >= 3.4 and width_px >= 250 and height_px <= 300 and ratio > best_ratio:
            best_blob = part.blob
            best_ratio = ratio

    return best_blob
