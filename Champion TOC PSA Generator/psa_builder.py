"""Professional Services Agreement generation.

Owns: loading the PSA template, populating fields, preserving formatting,
populating tables, applying signatures, and returning final document
bytes + warnings. Receives structured data from the Streamlit app; must
never import streamlit or contain UI code.

place_meghan_execution_signature() below is intentionally unused (dead
code, preserved verbatim from the original app.py) — the final "15.
SIGNATURES" page is deliberately left blank for both parties; Meghan only
signs the cover letter, and execution happens later via Adobe e-sign. Kept
here rather than deleted in case multi-consultant execution-page signing
is wanted in the future (see config.SIGNATURE_PROFILES).
"""

import io
import re
from copy import deepcopy
from pathlib import Path
from typing import Any, Dict, List, Sequence, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

from config import (
    ALL_SERVICES,
    DEFAULT_SIGNER_KEY,
    ENGAGEMENT_BOTH,
    ENGAGEMENT_HR,
    ENGAGEMENT_TA,
    SIGNERS,
    TEMPLATE_CANDIDATES,
    find_existing_file,
)
from scope_library import get_services_summary
from scoping import (
    agreement_label,
    current_agreement_date,
    apply_raas_defaults,
    build_replacements,
    clean,
    clean_multiline,
    combined_scope_from_data,
    engagement_includes_hr,
    engagement_includes_ta,
    first_name,
    format_rate,
    normalize_commitment_text,
    normalize_engagement,
    normalize_hours_display,
    normalize_hr_service,
    normalize_minimum_sentence,
    minimum_row_value,
    normalize_ta_service,
    tbc,
)
from word_helpers import (
    _clear_run_text_only,
    _copy_paragraph_format,
    _copy_run_format,
    _paragraph_contains_drawing,
    _run_contains_drawing,
    _set_row_minimum_height,
    clean_inserted_run,
    clear_cell_text_preserving_format,
    iter_all_paragraphs,
    remove_paragraph,
    remove_table_row,
    replace_everywhere,
    replace_text_in_paragraph,
    replace_visible_text_preserving_runs,
    set_checkbox,
    set_paragraph_value,
    set_run_black,
    set_table_cell,
)


def configure_service_cell(cell, selected_services: Sequence[str]) -> None:
    selected = {clean(service) for service in selected_services if clean(service)}
    for paragraph in cell.paragraphs:
        label = clean(paragraph.text)
        if label in ALL_SERVICES:
            set_checkbox(paragraph, label in selected)
        for run in paragraph.runs:
            set_run_black(run)

def set_service_cell_from_template(
    cell,
    engagement: str,
    hr_service: str,
    ta_service: str,
) -> None:
    """Retain the template's service-cell paragraph geometry and show only selected services."""
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        set_table_cell(cell, "To Be Confirmed")
        return

    # Capture the exact template paragraph XML before removing unused options.
    hr_heading = deepcopy(paragraphs[0]._p) if len(paragraphs) > 0 else None
    hr_option_map = {
        "HR Project Support": 1,
        "Fractional/Interim HR Support": 2,
        "HR Subscription": 3,
    }
    spacer = deepcopy(paragraphs[4]._p) if len(paragraphs) > 4 else None
    ta_heading = deepcopy(paragraphs[5]._p) if len(paragraphs) > 5 else None
    ta_option_map = {
        "Full Cycle Talent Acquisition Support": 6,
        "Sourcing Support": 7,
    }

    selected_xml = []
    selected_values: List[Tuple[str, bool]] = []

    if engagement_includes_hr(engagement) and hr_heading is not None:
        selected_xml.append(hr_heading)
        selected_values.append(("Human Resources Support:", True))
        option_index = hr_option_map.get(hr_service or "HR Project Support", 1)
        if option_index < len(paragraphs):
            selected_xml.append(deepcopy(paragraphs[option_index]._p))
            selected_values.append((f"☑ {hr_service or 'HR Project Support'}", False))

    if engagement_includes_ta(engagement) and ta_heading is not None:
        if selected_xml and spacer is not None:
            selected_xml.append(spacer)
            selected_values.append(("", False))
        selected_xml.append(ta_heading)
        selected_values.append(("Talent Acquisition Support:", True))
        option_index = ta_option_map.get(
            ta_service or "Full Cycle Talent Acquisition Support",
            6,
        )
        if option_index < len(paragraphs):
            selected_xml.append(deepcopy(paragraphs[option_index]._p))
            selected_values.append((
                f"☑ {ta_service or 'Full Cycle Talent Acquisition Support'}",
                False,
            ))

    if not selected_xml:
        selected_xml = [deepcopy(paragraphs[0]._p)]
        selected_values = [("To Be Confirmed", False)]

    for child in list(cell._tc):
        if child.tag == qn("w:p"):
            cell._tc.remove(child)

    for paragraph_xml in selected_xml:
        cell._tc.append(paragraph_xml)

    refreshed = list(cell.paragraphs)
    for paragraph, (text, bold) in zip(refreshed, selected_values):
        for content_control in list(paragraph._p.iter(qn("w:sdt"))):
            parent = content_control.getparent()
            if parent is not None:
                parent.remove(content_control)

        for text_node in paragraph._p.iter(qn("w:t")):
            text_node.text = re.sub(r"^[☐☑☒□■]+\s*", "", text_node.text or "")

        set_paragraph_value(paragraph, text, bold)

def blacken_red_template_text(doc: Document) -> None:
    for paragraph in iter_all_paragraphs(doc):
        for run in paragraph.runs:
            color = run.font.color.rgb
            if color and str(color).upper() in {"EE0000", "FF0000", "C00000"}:
                set_run_black(run)

def remove_duplicate_email_in_letter(doc: Document) -> None:
    """Remove repeated operations email text without rebuilding the paragraph."""
    target_email = "operations@tocgrp.com"

    for paragraph in doc.paragraphs:
        visible_text = "".join(
            node.text or ""
            for node in paragraph._p.iter(qn("w:t"))
        )

        if (
            "Please contact me or" not in visible_text
            or visible_text.lower().count(target_email) < 2
        ):
            continue

        seen_email = False

        for text_node in paragraph._p.iter(qn("w:t")):
            value = text_node.text or ""
            lowered = value.lower()
            search_from = 0
            pieces = []

            while True:
                index = lowered.find(target_email, search_from)
                if index < 0:
                    pieces.append(value[search_from:])
                    break

                pieces.append(value[search_from:index])

                if not seen_email:
                    pieces.append(value[index:index + len(target_email)])
                    seen_email = True

                search_from = index + len(target_email)

            text_node.text = "".join(pieces)

        current_text = "".join(
            node.text or ""
            for node in paragraph._p.iter(qn("w:t"))
        )
        current_text = re.sub(
            r"questions\.\s*\.",
            "questions.",
            current_text,
            flags=re.IGNORECASE,
        )

        if "questions.." in current_text:
            text_nodes = list(paragraph._p.iter(qn("w:t")))
            if text_nodes:
                text_nodes[0].text = current_text.replace("questions..", "questions.")
                for node in text_nodes[1:]:
                    node.text = ""
        break

def move_cover_letter_signature(doc: Document) -> None:
    """Anchor Meghan's signature to its own paragraph, not an absolute page spot.

    The template anchors the signature image at a fixed distance from the top
    of the PAGE. That only lines up under "Welcome to The O'Connor Group
    family!" when the letter body above happens to wrap to exactly the same
    number of lines as the template's placeholder text. In real Word, with
    real client text, the wrapping is very often slightly different, so the
    fixed page position ends up overlapping "Meghan Popoleo" underneath it
    (confirmed by a screenshot showing the name hidden behind the image).

    The fix is to anchor the image relative to the paragraph it already sits
    in ("Welcome to..."), offset by just that paragraph's own line height, so
    it always starts directly below that one line regardless of how much
    text precedes it in the letter.
    """
    for paragraph in doc.paragraphs:
        if "Welcome to The O’Connor Group family" not in clean(paragraph.text):
            continue

        for anchor in paragraph._p.iter(qn("wp:anchor")):
            position_v = anchor.find(qn("wp:positionV"))
            if position_v is None:
                continue

            position_v.set("relativeFrom", "paragraph")

            for child in list(position_v):
                position_v.remove(child)

            # This paragraph's own line height clears its text; the image
            # then flows below it, unaffected by anything above the
            # paragraph. Nudged down from the exact line height (12pt) to
            # 19pt per a side-by-side screenshot showing the template's own
            # gap here is a bit more generous than a bare line-height offset.
            offset = OxmlElement("wp:posOffset")
            offset.text = "241300"
            position_v.append(offset)

            anchor.set("behindDoc", "0")
            anchor.set("allowOverlap", "1")
            return

def swap_signer_signature_image(doc: Document, signer_config: Dict[str, Any]) -> List[str]:
    """Replace the cover-letter handwritten signature image for the selected signer.

    Keeps the template's existing image height (cy) exactly as-is and scales
    only the width (cx) to the new image's own aspect ratio, so different
    people's actual signatures don't get stretched or squished into a box
    shaped for someone else's. Vertical position is untouched here —
    move_cover_letter_signature owns that separately, so signer selection
    never affects layout beyond the image itself.

    Returns a list of warning strings; never raises. If the signature file
    is missing or the image cannot be located/loaded, the template's
    existing image is left in place rather than guessing.
    """
    warnings: List[str] = []
    image_path = signer_config.get("signature_path")
    name = signer_config.get("name", "the selected signer")

    if not image_path or not Path(image_path).exists():
        warnings.append(
            f"Signature image file for {name} was not found; the PSA was "
            "generated with the template's existing signature image instead."
        )
        return warnings

    for paragraph in doc.paragraphs:
        if "Welcome to The O’Connor Group family" not in clean(paragraph.text):
            continue

        blip = next(paragraph._p.iter(qn("a:blip")), None)
        extent = next(paragraph._p.iter(qn("wp:extent")), None)
        if blip is None or extent is None:
            warnings.append(
                f"Could not locate the cover-letter signature image to replace "
                f"it with {name}'s signature; the template's existing image "
                "was left in place."
            )
            return warnings

        old_cy = int(extent.get("cy"))

        # If the selected signer's image is byte-identical to what the
        # template already has embedded (true for the default signer today,
        # since her external signature file matches the template's own
        # embedded copy exactly), skip the swap entirely rather than
        # recomputing cx from pixel dimensions — that recomputation can
        # round to a few hundred EMU off the template's exact original
        # value, an unnecessary change with no formatting benefit.
        current_rid = blip.get(qn("r:embed"))
        current_rel = doc.part.rels.get(current_rid) if current_rid else None
        try:
            new_image_bytes = Path(image_path).read_bytes()
        except OSError:
            warnings.append(
                f"Could not load the signature image file for {name}; the "
                "PSA was generated with the template's existing signature "
                "image instead."
            )
            return warnings

        if current_rel is not None and current_rel.target_part.blob == new_image_bytes:
            return warnings

        try:
            new_rid, image = doc.part.get_or_add_image(str(image_path))
        except Exception:
            warnings.append(
                f"Could not load the signature image file for {name}; the "
                "PSA was generated with the template's existing signature "
                "image instead."
            )
            return warnings

        new_cx = int(old_cy * (image.px_width / image.px_height))
        blip.set(qn("r:embed"), new_rid)
        extent.set("cx", str(new_cx))
        for inner_ext in paragraph._p.iter(qn("a:ext")):
            inner_ext.set("cx", str(new_cx))
        return warnings

    warnings.append(
        f"Could not find the cover-letter paragraph to place {name}'s "
        "signature; the template's existing signature image was left in place."
    )
    return warnings

def rename_cover_page_byline(doc: Document, name: str) -> None:
    """Fix the front cover page's byline, which lives inside a text box, not
    a normal body paragraph.

    The cover page shows "THE O'CONNOR GROUP" / "MEGHAN POPOLEO" inside a
    Word text box, bound to the document's Author core property via a
    content control. doc.paragraphs only walks body-level paragraphs and
    never descends into text boxes, so the render_signer_to_document loop
    below silently missed this — every generated PSA kept showing "Meghan
    Popoleo" there regardless of the selected signer. Fixed two ways: (1)
    directly overwrite the visible cached text wherever it says exactly
    "Meghan Popoleo", found by walking the full XML tree (which, unlike
    doc.paragraphs, does descend into text boxes) so it displays correctly
    immediately, even in viewers that never recalculate Word fields; and
    (2) update the underlying Author property so the content control's
    bound value is correct too.
    """
    if not name:
        return

    doc.core_properties.author = name

    for t in doc.element.body.iter(qn("w:t")):
        if clean(t.text) == "Meghan Popoleo":
            t.text = name

def render_signer_to_document(doc: Document, signer_config: Dict[str, Any]) -> List[str]:
    """Populate every place the selected TOC signer's identity appears.

    Covers: the front cover page's byline (inside a text box — see
    rename_cover_page_byline), the cover-letter closing block (name +
    title; the "The O’Connor Group" company line stays the same for every
    signer), the handwritten signature image, and the final execution-page
    signature table's TOC side. Centralized here so signer selection never
    needs isolated patches scattered across the document-population code.
    """
    name = signer_config.get("name", "")
    title = signer_config.get("title", "")
    warnings: List[str] = []

    rename_cover_page_byline(doc, name)

    for paragraph in doc.paragraphs:
        if _paragraph_contains_drawing(paragraph):
            continue
        text = clean(paragraph.text)
        if text == "Meghan Popoleo" and name and text != name:
            set_paragraph_value(paragraph, name)
        elif text == "President" and title and text != title:
            set_paragraph_value(paragraph, title)

    warnings.extend(swap_signer_signature_image(doc, signer_config))

    signatures = find_signature_table(doc)
    if signatures is not None and len(signatures.rows) > 4 and len(signatures.columns) > 1:
        set_table_cell(signatures.cell(2, 1), f"Name: {name}")
        set_table_cell(signatures.cell(3, 1), f"Title: {title}")

    return warnings

def _replace_paragraph_with_plain_text(paragraph, value: str) -> None:
    """Rewrite a paragraph as plain text so hidden Word date fields cannot reappear."""
    value = "" if value is None else str(value)

    source_run = next(
        (run for run in paragraph.runs if not _run_contains_drawing(run)),
        None,
    )
    source_rpr = (
        deepcopy(source_run._r.rPr)
        if source_run is not None and source_run._r.rPr is not None
        else None
    )

    # Keep paragraph properties, but remove all text, fields, hyperlinks, and
    # content controls so Word cannot restore an old date when the file opens.
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)

    new_run = paragraph.add_run(value)
    if source_rpr is not None:
        current_rpr = new_run._r.get_or_add_rPr()
        new_run._r.remove(current_rpr)
        new_run._r.insert(0, source_rpr)
    clean_inserted_run(new_run)

def synchronize_agreement_dates(doc: Document) -> None:
    """Write one identical, field-free execution date everywhere."""
    agreement_date = current_agreement_date()

    for paragraph in iter_all_paragraphs(doc):
        text = clean(paragraph.text)
        if not text:
            continue

        if (
            'This Professional Services Agreement ("Agreement") is entered into as of'
            in text
        ):
            client_match = re.search(
                r'and between The O[\'’]Connor Group \(\"TOC\"\) and '
                r'(.+?) \(\"Client\"\)\.?$',
                text,
                flags=re.IGNORECASE,
            )
            client = client_match.group(1) if client_match else "Client"
            _replace_paragraph_with_plain_text(
                paragraph,
                f'This Professional Services Agreement ("Agreement") is entered into as of '
                f'{agreement_date} by and between The O\'Connor Group ("TOC") and '
                f'{client} ("Client").',
            )
            continue

        if (
            "The parties, intending to be legally bound" in text
            and "executed this Professional Services Agreement as of" in text
        ):
            _replace_paragraph_with_plain_text(
                paragraph,
                "The parties, intending to be legally bound, have executed this "
                f"Professional Services Agreement as of {agreement_date}.",
            )

def normalize_cover_signature_text(doc: Document) -> None:
    """Correct Meghan's typed cover-letter signature, including text boxes."""
    replacements = {
        "meghan popoleo": "Meghan Popoleo",
        "the o'connor group": "The O’Connor Group",
        "the o’connor group": "The O’Connor Group",
    }

    for paragraph in iter_all_paragraphs(doc):
        # The handwritten signature is anchored to a paragraph on the cover
        # page. Do not touch that paragraph through python-docx; the XML text
        # pass below can correct text-box capitalization without moving it.
        if _paragraph_contains_drawing(paragraph):
            continue

        text = clean(paragraph.text)
        replacement = replacements.get(text.lower())
        # Only rewrite when the casing is actually wrong. Rewriting text that
        # already matches forces the run through clean_inserted_run, which
        # turns the template's navy signature-block text black for no reason.
        if replacement and text != replacement:
            set_paragraph_value(paragraph, replacement)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

    # Floating text boxes are not exposed as python-docx Paragraph objects.
    roots = [doc._element]
    for section in doc.sections:
        roots.extend([section.header._element, section.footer._element])

    for root in roots:
        for text_node in root.iter(qn("w:t")):
            value = text_node.text or ""
            updated = re.sub(
                r"\bMeghan\s+Popoleo\b",
                "Meghan Popoleo",
                value,
                flags=re.IGNORECASE,
            )
            if updated.lower() in {"the o'connor group", "the o’connor group"}:
                updated = "The O’Connor Group"
            text_node.text = updated

def polish_travel_expenses_row(doc: Document) -> None:
    """Keep the travel row compact, readable, and on one page."""
    from docx.oxml import OxmlElement

    for table in doc.tables:
        for row in table.rows:
            if not row.cells or not any(
                "Travel Expenses" in clean(cell.text) for cell in row.cells
            ):
                continue

            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                tr_pr.append(OxmlElement("w:cantSplit"))

            # This is a minimum, not an exact height, so the text can still wrap.
            _set_row_minimum_height(row, 54)

            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
            return

def force_term_details_to_new_page(doc: Document) -> None:
    """Start Term Details on a new page without inserting a blank page."""
    first_table = find_term_details_table(doc)
    if first_table is None:
        return
    table_xml = first_table._tbl
    previous = table_xml.getprevious()

    # Remove legacy manual page-break runs from the blank paragraph before the
    # table. A standalone break can produce a blank page in some Word views.
    if previous is not None and previous.tag == qn("w:p"):
        for br in list(previous.iter(qn("w:br"))):
            if br.get(qn("w:type")) == "page":
                br.getparent().remove(br)

    # Put the break directly on the first table paragraph instead.
    first_paragraph = first_table.cell(0, 0).paragraphs[0]
    first_paragraph.paragraph_format.page_break_before = True

def compact_travel_expenses_row(doc: Document) -> None:
    """Preserve the template Travel Expenses row and prevent it from splitting."""
    from docx.oxml import OxmlElement

    for table in doc.tables:
        for row in table.rows:
            if not any("Travel Expenses" in cell.text for cell in row.cells):
                continue
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                tr_pr.append(OxmlElement("w:cantSplit"))
            return

def clear_billing_answer_cell(cell) -> None:
    """Completely clear a Billing Details answer cell, including hyperlinks."""
    for paragraph in cell.paragraphs:
        for hyperlink in list(paragraph._p.findall(qn("w:hyperlink"))):
            paragraph._p.remove(hyperlink)

        for text_node in paragraph._p.iter(qn("w:t")):
            text_node.text = ""

        for instruction in paragraph._p.iter(qn("w:instrText")):
            instruction.text = ""

        for simple_field in list(paragraph._p.findall(qn("w:fldSimple"))):
            paragraph._p.remove(simple_field)

        editable_runs = [
            run for run in paragraph.runs
            if not _run_contains_drawing(run)
        ]
        for run in editable_runs:
            _clear_run_text_only(run)

def polish_generated_tables(doc: Document) -> None:
    """Apply conservative table cleanup and leave Billing Details blank."""
    from docx.oxml import OxmlElement

    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                tr_pr.append(OxmlElement("w:cantSplit"))

    billing_labels = {
        "company name",
        "company name (dba)",
        "main phone number",
        "primary contact name & email",
        "billing contact",
        "billing contact name",
        "billing email",
        "billing address",
        "billing phone",
        "billing contact name & email",
        "full billing address for invoice records",
        "are you tax exempt? select yes or no.",
        "if your a/p department requires a po # on invoices, provide here:",
        "purchase order number",
        "po number",
        "accounts payable contact",
    }

    for table in doc.tables:
        first_column = [
            clean(row.cells[0].text).lower()
            for row in table.rows
            if row.cells
        ]

        label_matches = sum(
            1
            for cell_text in first_column
            if any(label in cell_text for label in billing_labels)
        )

        if label_matches < 2:
            continue

        for row in table.rows:
            if len(row.cells) > 1:
                clear_billing_answer_cell(row.cells[1])
        break

def fix_primary_scope_format(doc: Document) -> None:
    """Normalize inserted scope text and stop justified spacing from stretching it."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    primary_heading_seen = False

    for paragraph in doc.paragraphs:
        raw = clean_multiline(paragraph.text)

        if clean(paragraph.text) == "Primary Services":
            primary_heading_seen = True
            continue

        if not raw:
            continue

        if primary_heading_seen:
            normalized = re.sub(
                r"(\d+)\s*[-–]\s*(\d+)\s*hours?",
                r"\1–\2 hours",
                raw,
                flags=re.IGNORECASE,
            )
            normalized = re.sub(
                r"(\d+)\s*hours?\s+minimum",
                r"\1-hour minimum",
                normalized,
                flags=re.IGNORECASE,
            )
            set_paragraph_value(paragraph, normalized)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(10)
            paragraph.paragraph_format.line_spacing = 1.0
            # set_paragraph_value always forces inserted text to true black
            # via clean_inserted_run, but the template's own body text next
            # to this section (e.g. "Service Delivery Model" below) uses a
            # slate blue-gray, not black. Restore that same slate here so
            # the Scope of Services text matches the rest of the document's
            # body copy instead of standing out.
            for run in paragraph.runs:
                if not _run_contains_drawing(run):
                    run.font.color.rgb = RGBColor(0x49, 0x54, 0x67)
            return

        if raw.startswith("RaaS"):
            one_line = re.sub(r"\s+", " ", raw).strip()

            rate_match = re.search(
                r"\$\s*\d+(?:\.\d+)?\s*/\s*hour",
                one_line,
                flags=re.IGNORECASE,
            )
            minimum_match = re.search(
                r"(\d+(?:\.\d+)?)\s*[- ]?\s*hours?\s+minimum",
                one_line,
                flags=re.IGNORECASE,
            )
            estimate_match = re.search(
                r"estimated\s+time\s+to\s+fill\s*:?\s*"
                r"(\d+(?:\.\d+)?)\s*[-–]\s*(\d+(?:\.\d+)?)\s*hours?",
                one_line,
                flags=re.IGNORECASE,
            )

            lines = []
            rate = rate_match.group(0).replace(" ", "") if rate_match else ""
            if rate:
                lines.append(f"RaaS (Recruiting as a Service) — {rate}")
            else:
                lines.append("RaaS (Recruiting as a Service)")

            if minimum_match:
                lines.append(f"{minimum_match.group(1)}-hour minimum")

            if estimate_match:
                lines.append(
                    "Estimated time to fill: "
                    f"{estimate_match.group(1)}–{estimate_match.group(2)} hours"
                )

            # If the supplied scope contained additional content that was not
            # captured above, keep it rather than discarding it.
            if len(lines) == 1:
                lines = [
                    re.sub(
                        r"(\d+)\s*[-–]\s*(\d+)\s*hours?",
                        r"\1–\2 hours",
                        raw,
                        flags=re.IGNORECASE,
                    )
                ]

            set_paragraph_value(paragraph, "\n".join(lines))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.line_spacing = 1.0
            return

        if (
            raw.startswith("TOC will act as an extension")
            or raw.startswith("Human Resources Support")
            or raw.startswith("Talent Acquisition Support")
        ):
            normalized = re.sub(
                r"(\d+)\s*[-–]\s*(\d+)\s*hours?",
                r"\1–\2 hours",
                raw,
                flags=re.IGNORECASE,
            )
            set_paragraph_value(paragraph, normalized)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.line_spacing = 1.0
            return

def add_optional_estimated_hours_row(term_table, estimated: str) -> None:
    """Add an Estimated Hours row after Project Minimum only when supplied."""
    if not estimated:
        return

    # Do not add a duplicate row if the template or a prior pass already has one.
    for row in term_table.rows:
        if row.cells and clean(row.cells[0].text).lower() == "estimated hours":
            set_table_cell(row.cells[1], estimated)
            return

    # Clone the Project Minimum row so borders, shading, widths, and fonts match.
    source_index = None
    for index, row in enumerate(term_table.rows):
        if row.cells and clean(row.cells[0].text).lower() == "project minimum":
            source_index = index
            break

    if source_index is None:
        return

    source_row = term_table.rows[source_index]
    new_row_xml = deepcopy(source_row._tr)
    source_row._tr.addnext(new_row_xml)

    # Re-resolve rows after the XML insertion.
    new_row = term_table.rows[source_index + 1]
    set_table_cell(new_row.cells[0], "Estimated Hours", bold=True)
    set_table_cell(new_row.cells[1], estimated)

def hide_weekly_commitment_row_if_blank(
    term_table, engagement_type: str, weekly_commitment_raw: str
) -> None:
    """Hide the Agreement Summary's Weekly Commitment row for TA-only
    (RaaS) engagements when no weekly commitment was provided.

    For HR engagements (HR-only or combined HR+TA), Weekly Commitment is an
    important commercial term and always stays visible — unchanged from
    prior behavior, including its "To Be Confirmed" fallback. For a
    TA-only engagement, the weekly commitment is already described in the
    Scope of Services' Service Delivery Model, so an empty row here would
    only duplicate (or risk contradicting) that language. This only ever
    removes the one row; it never touches anything else in the table.
    """
    if engagement_includes_hr(engagement_type):
        return
    if not engagement_includes_ta(engagement_type):
        return
    if clean(weekly_commitment_raw):
        return

    for row in term_table.rows:
        if row.cells and clean(row.cells[0].text).lower() == "weekly commitment":
            remove_table_row(term_table, row)
            return

def find_term_details_table(doc: Document):
    """Locate the Term Details table by its labels instead of table position.

    Uses only labels that are always present in the Term Details table
    ("client", "service type", "hourly rate", "project minimum") plus
    "weekly commitment", which is intentionally hidden for some TA-only
    (RaaS) engagements (see hide_weekly_commitment_row_if_blank) — so
    detection must not depend on that one row being present.
    """
    required_labels = {
        "client",
        "service type",
        "hourly rate",
        "weekly commitment",
        "project minimum",
    }

    best_table = None
    best_score = 0

    for table in doc.tables:
        labels = {
            clean(row.cells[0].text).lower()
            for row in table.rows
            if row.cells
        }
        score = sum(
            1 for required in required_labels
            if any(required in label for label in labels)
        )
        if score > best_score:
            best_score = score
            best_table = table

    return best_table if best_score >= 3 else None

def find_standard_services_table(doc: Document):
    """Locate the Standard Services pricing table by its header text."""
    for table in doc.tables:
        if not table.rows:
            continue
        header_text = " | ".join(clean(cell.text).lower() for cell in table.rows[0].cells)
        if (
            "service" in header_text
            and ("rate" in header_text or "pricing" in header_text or "fee" in header_text)
            and ("notes" in header_text or "details" in header_text or "commitment" in header_text)
        ):
            return table
    return None

def find_signature_table(doc: Document):
    """Locate the execution-signature table even when the signature is an image."""
    for table in reversed(doc.tables):
        text = " ".join(
            clean(cell.text)
            for row in table.rows
            for cell in row.cells
        ).lower()
        signature_count = text.count("signature")
        has_name_rows = "name:" in text or "title:" in text or "date:" in text
        if signature_count >= 2 and has_name_rows and len(table.columns) >= 2:
            return table
    return None

def preserve_template_signature_position(doc: Document) -> None:
    """Leave Meghan's handwritten signature exactly where the template placed it."""
    # Intentionally no XML movement. The template is the source of truth for
    # the floating signature anchor; hard-coded page offsets can push it off-page.
    return

def merge_summary_and_term_details(doc: Document) -> None:
    """Keep Agreement Summary on its own page and place Term Details beneath it.

    The template's team artwork sits on a very tall transparent canvas and is
    what makes the summary page look empty. Remove only that artwork, preserve
    the page break before the Agreement Summary, and remove only the break that
    forces Term Details onto a separate page.
    """
    term_table = find_term_details_table(doc)
    if term_table is None:
        return

    # Lock the Agreement Summary to the beginning of its own page.
    summary_heading = next(
        (p for p in doc.paragraphs if clean(p.text) == "PROFESSIONAL SERVICES AGREEMENT"),
        None,
    )
    if summary_heading is not None:
        summary_heading.paragraph_format.page_break_before = True
        summary_heading.paragraph_format.keep_with_next = True

    # Note: the template's own images (including the team photo on the
    # preceding page) are left untouched here. An earlier version removed any
    # tall portrait image under the assumption it was an empty canvas hiding
    # behind the Agreement Summary page, but the only image matching that
    # heuristic is the team photo itself, so that step only deleted content
    # that belongs in the template.

    # Term Details should follow the summary on the same page.
    first_paragraph = term_table.cell(0, 0).paragraphs[0]
    first_paragraph.paragraph_format.page_break_before = False
    first_paragraph.paragraph_format.keep_with_next = True

    # Scope must still begin on a fresh page after Term Details.
    scope_heading = next(
        (p for p in doc.paragraphs if clean(p.text).startswith("1. SCOPE OF SERVICES")),
        None,
    )
    if scope_heading is not None:
        scope_heading.paragraph_format.page_break_before = True
        scope_heading.paragraph_format.keep_with_next = True

def place_meghan_execution_signature(doc: Document) -> None:
    """Left as a no-op: Meghan does not pre-sign the legal signature page.

    She signs only the cover letter (a courtesy/welcome signature). The
    contract's actual execution signatures, for both Client and TOC, are
    completed afterward through Adobe/e-signature once the parties agree to
    move forward, so this page must stay blank like the template's own
    "Signature" placeholder text on both sides.
    """
    return

def populate_template_specific_fields(doc: Document, data: Dict[str, str]) -> None:
    data = apply_raas_defaults(data)
    client = tbc(data.get("client_name"))
    contact = tbc(data.get("contact_name"))
    title = tbc(data.get("contact_title"))
    address_1 = tbc(data.get("address_1"))
    address_2 = clean(data.get("address_2"))
    email = clean(data.get("contact_email"))
    phone = clean(data.get("phone"))
    rate = format_rate(data.get("hourly_rate", ""))
    weekly = normalize_commitment_text(data.get("weekly_commitment", "")) or "To Be Confirmed"
    minimum = normalize_minimum_sentence(data.get("minimum_engagement", "")) or "To Be Confirmed"
    # Term Details' "Project Minimum" row label already says "Minimum" —
    # the value there should read "40 hours", not "Minimum 40 hours".
    # Every other use of `minimum` (e.g. the Standard Services pricing
    # table's Notes column) keeps the full "Minimum 40 hours" phrase,
    # since that column has no row label of its own to lean on.
    minimum_row_display = minimum_row_value(data.get("minimum_engagement", "")) or "To Be Confirmed"
    estimated = normalize_hours_display(data.get("estimated_hours", ""))
    engagement = normalize_engagement(
        data.get("engagement_type", ""),
        data.get("hr_service_type", ""),
        data.get("ta_service_type", ""),
    )
    hr_service = normalize_hr_service(data.get("hr_service_type", ""))
    ta_service = normalize_ta_service(data.get("ta_service_type", ""))

    # Cover letter placeholders.
    literal_replacements = {
        "NAME of Contact": contact,
        "Company Name.": client,
        "Address Line 1": address_1,
        "Address Line 2": address_2,
        "xxxxxx": client,
        "ENTER SCOPE OF WORK HERE": combined_scope_from_data(data),
        "$xx/hour": rate,
        "$xx": rate,
    }
    replace_everywhere(doc, literal_replacements)

    # Use the contact's first name in the salutation while retaining the full
    # legal name in the address block and signature section.
    greeting_name = first_name(contact) or contact
    for paragraph in doc.paragraphs:
        if clean(paragraph.text).lower().startswith("dear "):
            set_paragraph_value(paragraph, f"Dear {greeting_name},")
            break

    # Replace the standalone cover-letter title line only.
    for paragraph in doc.paragraphs:
        if clean(paragraph.text) == "Title":
            replace_visible_text_preserving_runs(paragraph, title)
            break

    # The generic words "Company Name" also appear as labels in billing and
    # signature tables, so only replace the standalone cover-letter paragraph.
    for paragraph in doc.paragraphs:
        if clean(paragraph.text) == "Company Name":
            set_paragraph_value(paragraph, client)

    # Replace the dated agreement summary even when the template date changes.
    for paragraph in iter_all_paragraphs(doc):
        if 'This Professional Services Agreement ("Agreement") is entered into as of' in paragraph.text:
            current_date = current_agreement_date()
            set_paragraph_value(
                paragraph,
                f'This Professional Services Agreement ("Agreement") is entered into as of '
                f'{current_date} by and between The O\'Connor Group ("TOC") and '
                f'{client} ("Client").',
            )

            for run in list(paragraph.runs):
                if (
                    any(True for _ in run._r.iter(qn("w:fldChar")))
                    or any(True for _ in run._r.iter(qn("w:instrText")))
                ):
                    paragraph._p.remove(run._r)

            for simple_field in list(paragraph._p.findall(qn("w:fldSimple"))):
                paragraph._p.remove(simple_field)

    # Populate the Services Summary in the existing template location.
    # The blank template paragraph has a serif style, so explicitly copy the
    # legal-summary paragraph formatting before writing the generated summary.
    #
    # Services Summary shows ONLY the high-level engagement category
    # ("Recruitment-as-a-Service (RaaS)" or "Human Resources Consulting") —
    # no pricing, hours, weekly commitment, or minimum. Those commercial
    # terms already live in the Agreement Summary / Term Details table
    # above this section; repeating them here would duplicate content the
    # approved template already shows elsewhere.
    summary_categories: List[str] = []
    if engagement_includes_ta(engagement) and ta_service:
        category = get_services_summary(ta_service)
        if category:
            summary_categories.append(category)
    if engagement_includes_hr(engagement) and hr_service:
        category = get_services_summary(hr_service)
        if category:
            summary_categories.append(category)

    agreement_summary_text = " & ".join(
        dict.fromkeys(summary_categories)
    ) or agreement_label(engagement)

    legal_summary = next(
        (
            paragraph
            for paragraph in iter_all_paragraphs(doc)
            if 'This Professional Services Agreement ("Agreement") is entered into as of'
            in paragraph.text
        ),
        None,
    )

    summary_paragraph = next(
        (
            paragraph
            for paragraph in iter_all_paragraphs(doc)
            if clean(paragraph.text).lower().startswith(
                ("services summary:", "service summary:")
            )
        ),
        None,
    )

    if summary_paragraph is None and legal_summary is not None:
        sibling = legal_summary._p.getnext()
        while sibling is not None and sibling.tag != qn("w:p"):
            sibling = sibling.getnext()
        if sibling is not None:
            candidate = Paragraph(sibling, legal_summary._parent)
            candidate_text = clean(candidate.text).lower()
            if not candidate_text or candidate_text.startswith(
                ("services summary", "service summary")
            ):
                summary_paragraph = candidate

    if summary_paragraph is not None and agreement_summary_text:
        source_run = next(
            (run for run in legal_summary.runs if not _run_contains_drawing(run)),
            None,
        ) if legal_summary is not None else None

        if legal_summary is not None:
            _copy_paragraph_format(legal_summary, summary_paragraph)
            try:
                summary_paragraph.style = legal_summary.style
            except (AttributeError, ValueError):
                pass

        for run in list(summary_paragraph.runs):
            if not _run_contains_drawing(run):
                _clear_run_text_only(run)

        label_run = summary_paragraph.runs[0] if summary_paragraph.runs else summary_paragraph.add_run("")
        label_run.text = "Services Summary: "
        value_run = summary_paragraph.add_run(agreement_summary_text)

        for run in (label_run, value_run):
            if source_run is not None:
                _copy_run_format(source_run, run)
            # Explicit fallback prevents the blank template slot from reverting
            # to Times New Roman or theme-green text in Word/preview renderers.
            run.font.name = source_run.font.name if source_run and source_run.font.name else "Arial"
            run.font.size = source_run.font.size if source_run and source_run.font.size else Pt(10)
            run._r.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), run.font.name)
            run._r.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), run.font.name)
            clean_inserted_run(run)

        label_run.bold = True
        value_run.bold = False
        summary_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        summary_paragraph.paragraph_format.space_before = Pt(8)
        summary_paragraph.paragraph_format.space_after = Pt(0)
        summary_paragraph.paragraph_format.line_spacing = 1.0
        summary_paragraph.paragraph_format.keep_with_next = False

    # Tailor the introductory letter so it does not mention a service that is
    # not part of the engagement.
    service_phrase = agreement_label(engagement)
    letter_service_phrase = {
        ENGAGEMENT_HR: "Human Resources",
        ENGAGEMENT_TA: "Talent Acquisition",
        ENGAGEMENT_BOTH: "Human Resources and Talent Acquisition",
    }.get(engagement, "professional services")
    # Use replace_text_in_paragraph (targeted, run-preserving substitution)
    # instead of rewriting the whole paragraph. This letter paragraph also
    # contains the italic "or" and the hyperlinked operations@tocgrp.com
    # address later in the same sentence; rewriting the entire paragraph text
    # collapsed all of that into one plain run, which duplicated the email
    # text and left remove_duplicate_email_in_letter() blanking out the real
    # hyperlink run, so the email lost its blue/underlined styling entirely.
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if "support your Human Resource and Recruitment-as-a-Service needs" in text:
            replace_text_in_paragraph(
                paragraph,
                "support your Human Resource and Recruitment-as-a-Service needs",
                f"support your {letter_service_phrase} needs",
            )
        elif "Human Resource and Recruitment-as-a-Service solutions" in text:
            replace_text_in_paragraph(
                paragraph,
                "Human Resource and Recruitment-as-a-Service solutions",
                f"{letter_service_phrase} solutions",
            )

    # Term Details table. Locate it by labels so Page 3 is populated even if
    # the template contains another table before it.
    term_table = find_term_details_table(doc)
    if term_table is not None:
        client_address = "\n".join(
            value for value in [client, address_1, address_2] if clean(value)
        )
        set_table_cell(term_table.cell(1, 1), client_address)

        service_lines: List[Tuple[str, bool]] = []

        if engagement_includes_hr(engagement):
            service_lines.extend(
                [
                    ("Human Resources Support:", True),
                    (f"☑ {hr_service or 'HR Project Support'}", False),
                ]
            )

        if engagement_includes_ta(engagement):
            if service_lines:
                service_lines.append((" ", False))

            service_lines.extend(
                [
                    ("Talent Acquisition Support:", True),
                    (
                        f"☑ {ta_service or 'Full Cycle Talent Acquisition Support'}",
                        False,
                    ),
                ]
            )

        service_cell = term_table.cell(4, 1)
        set_service_cell_from_template(
            service_cell,
            engagement,
            hr_service,
            ta_service,
        )

        set_table_cell(
            term_table.cell(5, 1),
            f"{rate} (30-minute increments)",
        )

        set_table_cell(term_table.cell(6, 1), weekly)
        set_table_cell(term_table.cell(7, 1), minimum_row_display)
        add_optional_estimated_hours_row(term_table, estimated)
        hide_weekly_commitment_row_if_blank(
            term_table, engagement, data.get("weekly_commitment", "")
        )

    # Standard Services table: keep only the services in this agreement.
    standard_table = find_standard_services_table(doc)
    if standard_table is not None:

        template_row_xml = (
            deepcopy(standard_table.rows[1]._tr)
            if len(standard_table.rows) > 1
            else None
        )

        while len(standard_table.rows) > 1:
            remove_table_row(
                standard_table,
                standard_table.rows[-1],
            )

        note_parts: List[str] = []
        if estimated:
            note_parts.append(f"Estimated {estimated} per search")
        if weekly != "To Be Confirmed":
            note_parts.append(weekly)
        if minimum != "To Be Confirmed":
            note_parts.append(minimum)
        note = "; ".join(note_parts) or "To Be Confirmed"

        service_rows: List[Tuple[str, str, str]] = []

        if engagement_includes_hr(engagement):
            service_rows.append(
                (
                    hr_service or "HR Project Support",
                    rate,
                    note,
                )
            )

        if engagement_includes_ta(engagement):
            service_rows.append(
                (
                    ta_service or "Full Cycle Talent Acquisition Support",
                    rate,
                    note,
                )
            )

        for service_name, service_rate, service_note in service_rows:
            if template_row_xml is not None:
                standard_table._tbl.append(deepcopy(template_row_xml))
                cells = standard_table.rows[-1].cells
            else:
                cells = standard_table.add_row().cells

            set_table_cell(cells[0], service_name, True)
            set_table_cell(cells[1], service_rate)
            set_table_cell(cells[2], service_note)

    # Remove irrelevant minimum-engagement boilerplate and replace retained
    # language with the actual deal terms.
    paragraphs = list(doc.paragraphs)

    for index, paragraph in enumerate(paragraphs):
        text = clean(paragraph.text)

        if text == "(Recruitment-as-a-Service)":
            if not engagement_includes_ta(engagement):
                remove_paragraph(paragraph)
                if index + 1 < len(paragraphs):
                    remove_paragraph(paragraphs[index + 1])
            elif index + 1 < len(paragraphs):
                if minimum:
                    minimum_for_sentence = re.sub(
                        r"^Minimum\s+",
                        "",
                        minimum,
                        flags=re.IGNORECASE,
                    )
                    set_paragraph_value(
                        paragraphs[index + 1],
                        (
                            "This Talent Acquisition engagement is subject to a "
                            f"minimum commitment of {minimum_for_sentence}."
                        ),
                    )
                else:
                    set_paragraph_value(
                        paragraphs[index + 1],
                        (
                            "Any minimum engagement requirement will be confirmed "
                            "in writing before services begin."
                        ),
                    )

        elif text in {
            "(Fractional/ Interim Human Resources Services)",
            "(Fractional/Interim Human Resources Services)",
        }:
            if not engagement_includes_hr(engagement):
                remove_paragraph(paragraph)
                if index + 1 < len(paragraphs):
                    remove_paragraph(paragraphs[index + 1])
            elif index + 1 < len(paragraphs):
                set_paragraph_value(
                    paragraph,
                    "(Fractional/Interim Human Resources Services)",
                )

                minimum_for_sentence = re.sub(
                    r"^Minimum\s+",
                    "",
                    minimum,
                    flags=re.IGNORECASE,
                )
                if minimum == "To Be Confirmed":
                    sentence = "a minimum commitment to be confirmed"
                else:
                    sentence = f"a minimum commitment of {minimum_for_sentence}"

                set_paragraph_value(
                    paragraphs[index + 1],
                    f"This engagement requires {sentence}.",
                )

    # Remove clearly TA-only legal bullets from HR-only agreements.
    if engagement == ENGAGEMENT_HR:
        ta_only_starts = [
            "Client agrees to provide timely interview feedback",
            "Fees are payable regardless of hiring outcomes",
            "TOC does not guarantee candidate placement",
        ]

        for paragraph in list(doc.paragraphs):
            if any(
                clean(paragraph.text).startswith(prefix)
                for prefix in ta_only_starts
            ):
                remove_paragraph(paragraph)

    # Make the TOC responsibility match the selected service. Use a targeted
    # substring replacement (not a whole-paragraph rewrite) so the bullet's
    # own numbering/list formatting is undisturbed, then restore the
    # template's slate body-text color: clean_inserted_run always forces
    # black, but this bullet (unlike a red placeholder) was already the
    # correct slate color in the template, so black would stand out here.
    for paragraph in doc.paragraphs:
        if (
            clean(paragraph.text)
            == "Provide professional Human Resource Services and/or Recruitment Services"
        ):
            replace_text_in_paragraph(
                paragraph,
                "Human Resource Services and/or Recruitment Services",
                service_phrase,
            )
            for run in paragraph.runs:
                if not _run_contains_drawing(run):
                    run.font.color.rgb = RGBColor(0x49, 0x54, 0x67)
            break

    # Billing Details are cleared by polish_generated_tables(), which locates
    # the section by its labels instead of relying on a fragile table index.

    # Signature table. Locate it by structure so it still works when Meghan's
    # handwritten signature is stored only as an image.
    signatures = find_signature_table(doc)
    if signatures is not None and len(signatures.rows) > 4 and len(signatures.columns) > 1:
        set_table_cell(signatures.cell(0, 0), client)
        set_table_cell(signatures.cell(0, 1), "The O'Connor Group")
        set_table_cell(signatures.cell(2, 0), f"Name: {contact}")
        set_table_cell(signatures.cell(3, 0), f"Title: {title}")
        set_table_cell(signatures.cell(4, 0), "Date:")
        # Signer name/title (cells 2,1 / 3,1) are populated by
        # render_signer_to_document, called from build_psa, so there is one
        # place that sets the TOC signer's identity rather than a hardcoded
        # name duplicated here.
        set_table_cell(signatures.cell(4, 1), "Date:")
        clear_cell_text_preserving_format(signatures.cell(1, 1))
        set_table_cell(signatures.cell(1, 1), "Signature")
        # set_table_cell always forces black via clean_inserted_run, but the
        # template's own "Signature" label (both sides) is slate, not black.
        for paragraph in signatures.cell(1, 1).paragraphs:
            for run in paragraph.runs:
                if not _run_contains_drawing(run):
                    run.font.color.rgb = RGBColor(0x49, 0x54, 0x67)

    remove_duplicate_email_in_letter(doc)
    normalize_cover_signature_text(doc)
    synchronize_agreement_dates(doc)
    merge_summary_and_term_details(doc)

def keep_signature_section_together(doc: Document) -> None:
    """Keep the signature heading with its table without forcing the whole ending together."""
    from docx.oxml import OxmlElement

    heading = next(
        (p for p in doc.paragraphs if clean(p.text).upper().startswith("15. SIGNATURES")),
        None,
    )
    if heading is not None:
        heading.paragraph_format.page_break_before = True
        heading.paragraph_format.keep_with_next = True

    # Find the signature table by structure (find_signature_table), not by a
    # hardcoded signer name — that name changes per the selected TOC signer.
    signature_table = find_signature_table(doc)
    if signature_table is not None:
        for row in signature_table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                tr_pr.append(OxmlElement("w:cantSplit"))

def remove_accidental_blank_break_paragraphs(doc: Document) -> None:
    """Preserve template-authored page breaks and section pagination."""
    return

def keep_narrative_paragraphs_together(doc: Document) -> None:
    """Stop the "what's next" intro paragraphs from splitting mid-sentence.

    The template has no explicit page break tying the cover letter to the
    following page, so depending on how long the letter body runs for a
    given client, these paragraphs can land right at the page boundary and
    split awkwardly (confirmed by a screenshot: a sentence ending "...gain a
    competitive edge" on one page and continuing "where it matters most,
    people." on the next). Setting keep_together means Word will move the
    whole paragraph to the next page instead of breaking inside it.
    """
    targets = (
        "At The O’Connor Group, we partner with",
        "Acting as a seamless extension of your",
        "What sets us apart is our relationship",
    )
    for paragraph in doc.paragraphs:
        text = clean(paragraph.text)
        if any(text.startswith(prefix) for prefix in targets):
            paragraph.paragraph_format.keep_together = True

def apply_final_formatting(doc: Document) -> None:
    """Apply one conservative cleanup pass.

    The template remains responsible for fonts, spacing, page layout, colors,
    and the floating cover-letter signature. No global paragraph rebuilding
    occurs, and no blanket recoloring pass runs here: every function that
    inserts client data already calls clean_inserted_run/set_paragraph_value
    on exactly the runs it writes, which is what keeps generated text black.
    A document-wide sweep would also repaint the template's own untouched
    navy headings and slate body text, which is not desired.
    """
    fix_primary_scope_format(doc)
    polish_generated_tables(doc)
    polish_travel_expenses_row(doc)
    keep_narrative_paragraphs_together(doc)
    synchronize_agreement_dates(doc)
    normalize_cover_signature_text(doc)
    move_cover_letter_signature(doc)
    remove_accidental_blank_break_paragraphs(doc)
    keep_signature_section_together(doc)

    # Preserve the template-authored handwritten signature anchor.
    preserve_template_signature_position(doc)

def force_field_recalculation_on_open(doc: Document) -> None:
    """Make Word recompute every field (page numbers, dates, etc.) as soon
    as the generated PSA is opened.

    python-docx never recalculates Word fields — a document built this way
    keeps whatever page-number/date result was last cached in the template
    when it was saved in real Word, so every generated PSA can show a
    stale "Page | 1" on every page until someone manually refreshes fields
    (Ctrl+A, F9) or the app that opens it happens to recalculate on its
    own. Setting <w:updateFields/> in settings.xml tells Word to refresh
    all fields automatically the moment the document opens, so this never
    depends on the viewer. Safe to add unconditionally; it only affects
    field recalculation timing, nothing about layout or formatting.
    """
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is not None:
        existing.set(qn("w:val"), "true")
        return

    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.insert(0, update_fields)

def build_psa(data: Dict[str, str]) -> Tuple[bytes, List[str]]:
    warnings: List[str] = []

    template_path = find_existing_file(TEMPLATE_CANDIDATES)
    if template_path is None:
        raise FileNotFoundError(
            "No PSA template was found. Add 'Meghan PSA Template.docx' or "
            "'Meghan PSA.docx' to the same folder as app.py."
        )

    if not clean(data.get("client_name")):
        raise ValueError("Client Name is required.")

    if not clean(data.get("hourly_rate")):
        raise ValueError("Pricing or Hourly Rate is required.")

    engagement_type = normalize_engagement(
        data.get("engagement_type", ""),
        data.get("hr_service_type", ""),
        data.get("ta_service_type", ""),
    )

    if not engagement_type:
        raise ValueError(
            "Select an Engagement Type before generating the PSA."
        )

    if (
        engagement_includes_hr(engagement_type)
        and not clean(data.get("hr_scope_of_work"))
    ):
        warnings.append(
            "The HR Scope of Work field was left blank; the approved TOC "
            "Scope Library baseline scope was used instead."
        )

    if (
        engagement_includes_ta(engagement_type)
        and not clean(data.get("ta_scope_of_work"))
    ):
        warnings.append(
            "The Talent Acquisition Scope of Work field was left blank; "
            "the approved TOC Scope Library baseline scope was used instead."
        )

    doc = Document(str(template_path))

    replace_everywhere(
        doc,
        build_replacements(data),
    )

    populate_template_specific_fields(
        doc,
        data,
    )

    signer_key = clean(data.get("signer")) or DEFAULT_SIGNER_KEY
    signer_config = SIGNERS.get(signer_key)
    if signer_config is None:
        warnings.append(
            f"'{signer_key}' is not a recognized TOC signer; used the default "
            f"signer ({DEFAULT_SIGNER_KEY}) instead."
        )
        signer_config = SIGNERS[DEFAULT_SIGNER_KEY]
    warnings.extend(render_signer_to_document(doc, signer_config))

    apply_final_formatting(doc)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return output.getvalue(), warnings
